import docx
import os
import logging
import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
import re
import shutil
import glob

# Initialize the saved word list as empty
saved_word_list = []

# Set the initial working directory
initial_working_directory = r'C:\Users\HamburgerC\Desktop\Bracketing code'

# Variable to track whether processing should occur
processing_enabled = False

# Output folder for processed files
output_folder = r'C:\Users\HamburgerC\Desktop\Bracketing code'  # Hardcoded output folder path

# Function to select a folder
def select_folder():
    folder_path = filedialog.askdirectory(initialdir=initial_working_directory)
    if folder_path:
        process_folder(folder_path)

# Function to select a Word document
def select_document():
    doc_path = filedialog.askopenfilename(initialdir=initial_working_directory, filetypes=[("Word Files", "*.docx")])
    if doc_path:
        process_document(doc_path)

# Function to add double brackets to words in text (with a maximum of two brackets)
def add_double_brackets_to_words_in_text(input_text, word_list):
    result_text = input_text

    for word in word_list:
        # Check if the word already contains double brackets
        if f"[[{word}]]" not in result_text:
            # Use a case-sensitive replacement
            result_text = result_text.replace(word, f"[[{word}]]")
        else:
            # If double brackets are already present, replace any occurrences of triple brackets with double brackets
            result_text = result_text.replace(f"[[[{word}]]]", f"[[{word}]]")

    return result_text

# Function to extract text from a Word document
def extract_text_from_docx(docx_filename):
    try:
        doc = docx.Document(docx_filename)
        text = '\n'.join([paragraph.text.strip() for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        logging.error(f"Error: Unable to read the Word document. Error: {str(e)}")
        return ""

# Function to clean up extra brackets beside the two that should be there
def clean_up_extra_brackets(input_text):
    cleaned_text = input_text

    # Use a regular expression to remove extra brackets beside the two that should be there
    cleaned_text = re.sub(r'\[\[+([^\[\]]+)\]\]+', r'[[\1]]', cleaned_text)

    return cleaned_text

# Function to process a Word document
def process_word_document(input_filename, word_list, result_filename, num_iterations=3):
    try:
        # Load the Word document
        doc_text = extract_text_from_docx(input_filename)

        # Initialize an empty list to store the processed paragraphs
        processed_paragraphs = []

        for _ in range(num_iterations):
            # Process the document text
            processed_text = add_double_brackets_to_words_in_text(doc_text, word_list)
            processed_text = clean_up_extra_brackets(processed_text)
            processed_paragraphs.append(processed_text)

        result_text = '\n'.join(processed_paragraphs)

        # Save the result to the specified Word document without adding "_processed"
        full_output_path = os.path.abspath(result_filename)

        # Check if the file already exists and prompt for overwrite
        if os.path.exists(full_output_path):
            confirm = messagebox.askyesno("File Exists", f"The file '{full_output_path}' already exists. Do you want to overwrite it?")
            if not confirm:
                logging.error("Processing canceled.")
                return False

        # Create a new Word document to store the result
        result_doc = docx.Document()
        result_doc.add_paragraph(result_text)

        # Save the result to the specified Word document
        result_doc.save(full_output_path)

        return True
    except Exception as e:
        logging.error(f"Error: An error occurred while processing the document. Error: {str(e)}")
        return False

# Function to check for and limit the number of brackets around words (maximum of two brackets)
def limit_brackets(word_list):
    for i in range(len(word_list)):
        word = word_list[i]
        # Use a while loop to remove extra brackets until only a maximum of two brackets remain
        while "[[" in word:
            word = word.replace("[[", "[")
        while "]]" in word:
            word = word.replace("]]", "]")
        word_list[i] = word

# Function to process an individual Word document
def process_document(doc_path):
    try:
        global processing_enabled  # Access the global variable

        if processing_enabled:  # Check if processing is enabled
            # Create a sanitized output file name by replacing problematic characters
            output_filename = re.sub(r'[\/:*?"<>|]', '_', doc_path)
            output_filename = os.path.join(output_folder, output_filename + ".docx")

            # Process the Word document and display a message based on success or failure
            if process_word_document(doc_path, saved_word_list, output_filename):
                messagebox.showinfo("Success", f"Processing completed successfully. Result saved as '{output_filename}'")
            else:
                messagebox.showerror("Error", "Processing failed.")
    except Exception as e:
        logging.error(f"Error: An error occurred while processing the document. Error: {str(e)}")

# Function to process a file and double bracket its content
def process_file(file_path):
    try:
        global processing_enabled  # Access the global variable

        if processing_enabled:  # Check if processing is enabled
            # Read the file content
            with open(file_path, 'r', encoding='utf-8') as file_content:
                file_text = file_content.read()

            # Create a sanitized output file name by replacing problematic characters
            output_filename = re.sub(r'[\/:*?"<>|]', '_', os.path.basename(file_path))
            output_filename = os.path.join(output_folder, output_filename + ".txt")

            # Double bracket the content and save it to the specified file
            result_text = add_double_brackets_to_words_in_text(file_text, saved_word_list)

            # Save the result to the specified file
            with open(output_filename, 'w', encoding='utf-8') as result_file:
                result_file.write(result_text)

            messagebox.showinfo("Success", f"Processing completed successfully. Result saved as '{output_filename}'")
    except Exception as e:
        logging.error(f"Error: An error occurred while processing the file. Error: {str(e)}")

# Function to update word list
def update_word_list(text_box):
    global saved_word_list, processing_enabled
    saved_word_list = text_box.get("1.0", "end-1c").splitlines()
    # Check and limit the number of brackets around words
    limit_brackets(saved_word_list)
    
    # Enable processing if there are words in the list
    if saved_word_list:
        processing_enabled = True
    else:
        processing_enabled = False
        messagebox.showwarning("Warning", "Word list is empty. Processing is disabled.")

# Function to update the text box with the word list
def update_text_box(word_list, text_box):
    text_box.delete("1.0", tk.END)
    for word in word_list:
        text_box.insert(tk.END, word + "\n")

# Function to save the word list to a text file
def save_word_list(text_box):
    global saved_word_list
    file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt")], initialdir=initial_working_directory)
    if file_path:
        with open(file_path, "w") as file:
            for word in saved_word_list:
                file.write(word + "\n")

# Function to load the word list from a text file
def load_word_list(text_box):
    global saved_word_list
    file_path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt")], initialdir=initial_working_directory)
    if file_path:
        with open(file_path, "r") as file:
            saved_word_list = [line.strip() for line in file]
            # Check and limit the number of brackets around words
            limit_brackets(saved_word_list)
            update_text_box(saved_word_list, text_box)  # Update the text box with the loaded word list

# Function to start the GUI application
def start_gui():
    # Create the main application window
    app = tk.Tk()
    app.title("Markdown to Word Document Processor")

    # Create and configure widgets
    text_box = tk.Text(app, height=10, width=40, font=("Arial", 12))
    text_box.pack(padx=20, pady=10)

    button_frame = tk.Frame(app)
    button_frame.pack(pady=20)

    # Buttons to select a folder
    select_folder_button = tk.Button(button_frame, text="Select Folder", font=("Arial", 14), command=select_folder)
    select_folder_button.grid(row=0, column=0, padx=5)

    # Buttons to select a Word document
    select_doc_button = tk.Button(button_frame, text="Select Document", font=("Arial", 14), command=select_document)
    select_doc_button.grid(row=0, column=1, padx=5)

    update_word_list_button = tk.Button(app, text="Update Word List", font=("Arial", 14), command=lambda: update_word_list(text_box))
    update_word_list_button.pack(pady=10)

    save_word_list_button = tk.Button(app, text="Save Word List", font=("Arial", 14), command=lambda: save_word_list(text_box))
    save_word_list_button.pack(pady=10)

    load_word_list_button = tk.Button(app, text="Load Word List", font=("Arial", 14), command=lambda: load_word_list(text_box))
    load_word_list_button.pack(pady=10)

    try:
        app.mainloop()
    except Exception as e:
        print(f"An error occurred: {str(e)}")

# Call the start_gui function when the script is executed
if __name__ == "__main__":
    start_gui()
