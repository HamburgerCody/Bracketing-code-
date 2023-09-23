import docx
import os
import logging

# Define constants for file paths
DOCX_FILENAME = r"C:\Users\HamburgerC\Desktop\Bracketing code\Unbracketed.docx"
RESULT_FILENAME = r"C:\Users\HamburgerC\Desktop\Bracketing code\Processed_Bracketed.docx"

# Configure logging without creating a log file
logging.basicConfig(level=logging.ERROR)

def add_double_brackets_to_words_in_text(input_text, word_list):
    words = input_text.split()
    result = []

    for word in words:
        # Check if the word matches any word in the provided list (case-insensitive)
        if any(word.lower() == item.lower() for item in word_list):
            result.append(f"[[{word}]]")  # Double brackets
        else:
            result.append(word)

    result_text = ' '.join(result)
    return result_text

def extract_text_from_docx(docx_filename):
    try:
        doc = docx.Document(docx_filename)
        text = '\n'.join([paragraph.text.strip() for paragraph in doc.paragraphs])
        return text
    except FileNotFoundError as e:
        logging.error(f"Error: The file '{docx_filename}' was not found. Error: {str(e)}")
        return ""

def process_word_document(docx_filename, word_list, result_filename):
    try:
        # Load the Word document
        doc_text = extract_text_from_docx(docx_filename)
    except FileNotFoundError as e:
        logging.error(f"Error: The file '{docx_filename}' was not found. Error: {str(e)}")
        return False

    # Split the document text into paragraphs
    paragraphs = doc_text.split('\n')

    # Initialize an empty list to store the processed paragraphs
    processed_paragraphs = []

    # Process each paragraph
    for paragraph in paragraphs:
        processed_paragraph = add_double_brackets_to_words_in_text(paragraph, word_list)
        processed_paragraphs.append(processed_paragraph)

    # Combine the processed paragraphs into a single text
    result_text = '\n'.join(processed_paragraphs)

    # Create a new Word document to store the result
    result_doc = docx.Document()
    result_doc.add_paragraph(result_text)

    # Save the result to a new Word document
    result_doc.save(result_filename)
    return True

if __name__ == "__main__":
    while True:
        # Prompt the user for the path to the Word List Word document
        word_list_path = input("Enter the path to the Word List Word document (or 'q' to quit): ")

        if word_list_path.lower() == 'q':
            break

        # Read the word list from the specified Word document
        word_list_text = extract_text_from_docx(word_list_path)

        if not word_list_text:
            print("Word list is empty or could not be read.")
        else:
            # Split the word list text into individual words
            word_list = word_list_text.split()

            # Print the word list for debugging
            print("Word List:", word_list)

            # Process the Word document and check if it was successful
            if process_word_document(DOCX_FILENAME, word_list, RESULT_FILENAME):
                print(f"Processing completed successfully. Result saved as '{RESULT_FILENAME}'")
            else:
                print("Processing failed.")
