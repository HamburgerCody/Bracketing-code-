import docx
import os
import logging
import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
from tkinter import simpledialog

logging.basicConfig(level=logging.ERROR)

# ... Other functions ...

def process_word_document(docx_filename, word_list, result_filename):
    try:
        doc = docx.Document(docx_filename)
    except Exception as e:
        logging.error(f"Error: Unable to read the Word document. Error: {str(e)}")
        return False

    # Create a new document to store the result
    result_doc = docx.Document()

    for paragraph in doc.paragraphs:
        processed_paragraph = add_double_brackets_to_words_in_text(paragraph.text, word_list)
        result_doc.add_paragraph(processed_paragraph)

    try:
        # Save the result to the specified Word document
        result_doc.save(result_filename)
        return True
    except Exception as e:
        logging.error(f"Error: Unable to save the Word document. Error: {str(e)}")
        return False

# ... Rest of the code ...
