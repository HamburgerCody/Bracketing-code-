import docx
import os

def read_word_list_from_file(file_path):
    try:
        with open(file_path, 'r') as file:
            word_list = [line.strip() for line in file.readlines()]
            return word_list
    except FileNotFoundError:
        print(f"Error: The file '{file_path}' was not found.")
        return []

def bracket_matching_words(input_text, word_list):
    words = input_text.split()
    result = []

    for word in words:
        # Check if the word matches any word in the provided list (case-insensitive)
        if any(word.lower() == item.lower() for item in word_list):
            result.append(f"[{word}]")
        else:
            result.append(word)

    result_text = ' '.join(result)
    return result_text

def process_word_document(docx_filename, word_list, result_filename):
    try:
        # Load the Word document
        doc = docx.Document(docx_filename)
    except FileNotFoundError:
        print(f"Error: The file '{docx_filename}' was not found.")
        return

    # Initialize an empty list to store the paragraphs
    paragraphs = []

    # Iterate through the paragraphs in the document
    for paragraph in doc.paragraphs:
        # Process each paragraph's text and add it to the list
        processed_paragraph = bracket_matching_words(paragraph.text, word_list)
        paragraphs.append(processed_paragraph)

    # Create a new Word document to store the result
    result_doc = docx.Document()

    # Add processed paragraphs to the result document
    for paragraph in paragraphs:
        result_doc.add_paragraph(paragraph)

    # Save the result to a new Word document
    try:
        result_doc.save(result_filename)
        print(f"Processing successful. Processed document saved as '{result_filename}'.")
    except Exception as e:
        print(f"Error: An error occurred while saving the processed document: {str(e)}")
        print("Processing failed.")

if __name__ == "__main__":
    # Prompt the user to change the path of the word list
    change_word_list_path = input("Do you want to change the path of the word list? (Yes/No): ").strip().lower()
    
    if change_word_list_path == "yes":
        # Ask the user to input the new path for the word list
        new_word_list_path = input("Enter the new path for the word list: ").strip()
        
        # Check if the new path is valid
        if os.path.exists(new_word_list_path):
            word_list = read_word_list_from_file(new_word_list_path)
        else:
            print("Error: The specified path does not exist. Using the default path.")
            word_list = read_word_list_from_file(r"C:\Users\HamburgerC\Desktop\Bracketing code\Word List\Finley Thornwood\Word_List.docx")
    else:
        # Use the default path for the word list
        word_list = read_word_list_from_file(r"C:\Users\HamburgerC\Desktop\Bracketing code\Word List\Finley Thornwood\Word_List.docx")
    
    word_set = set(word.lower() for word in word_list)

    # Specify the new path for the input Word document
    docx_filename = r"C:\Users\HamburgerC\Desktop\Bracketing code\Unbracketed.docx"
    
    # Specify the new path for the output processed Word document
    result_filename = r"C:\Users\HamburgerC\Desktop\Bracketing code\Processed_Bracketed.docx"

    # Call the function to process the Word document
    process_word_document(docx_filename, word_set, result_filename)
