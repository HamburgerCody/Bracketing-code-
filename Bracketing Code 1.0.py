import docx

def bracket_matching_words(input_text, word_set):
    words = input_text.split()
    result = []

    for word in words:
        # Check if the word matches any word in the provided set (case-insensitive)
        if word.lower() in word_set:
            result.append(f"[{word}]")
        else:
            result.append(word)

    result_text = ' '.join(result)
    return result_text

def process_word_document(docx_filename, word_set, result_filename):
    try:
        # Load the Word document
        doc = docx.Document(docx_filename)
    except FileNotFoundError:
        print(f"Error: The file '{docx_filename}' was not found.")
        return

    # Initialize an empty list to store the paragraphs
    paragraphs = []

    # Iterate through the paragraphs in the document
    for paragraph in doc.paragraphs:
        # Process each paragraph's text and add it to the list
        processed_paragraph = bracket_matching_words(paragraph.text, word_set)
        paragraphs.append(processed_paragraph)

    # Create a new Word document to store the result
    result_doc = docx.Document()

    # Add processed paragraphs to the result document
    for paragraph in paragraphs:
        result_doc.add_paragraph(paragraph)

    # Save the result to a new Word document
    result_doc.save(result_filename)

    print(f"Processed document saved as '{result_filename}'")

if __name__ == "__main__":
    # List of words to match
    word_list = [
        "Finley Thornwood",
        "Spell Scenario",
        "Abeir-Toril",
        "Faerûn",
        "Moonshadow Rangers",
        "The Greenwardens",
        "Aric Stonebrook",
        "Druidcraft",
        "Guidance",
        "Cure Wounds",
        "Appearance",
        "Half-Elf",
        "Druid",
        "Common",
        "Druidic",
        "Elvish",
        "Sylvan"
    ]

    # Convert the word list to lowercase for case-insensitive matching
    word_set = set(word.lower() for word in word_list)

    # Specify the path of the Word document to process and the result filename
    docx_filename = r"C:\Users\HamburgerC\Desktop\Bracketing code\Temp.docx"
    result_filename = r"C:\Users\HamburgerC\Desktop\Bracketing code\Processed_Temp.docx"

    # Call the function to process the Word document
    process_word_document(docx_filename, word_set, result_filename)
