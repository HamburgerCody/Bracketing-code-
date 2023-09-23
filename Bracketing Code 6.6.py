import docx
import os
import logging
import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
from tkinter import simpledialog  # For user input

# Configure logging without creating a log file
logging.basicConfig(level=logging.ERROR)

# Initialize the saved word list as empty
saved_word_list = []

# Function to process the document
def process_document():
    global saved_word_list
    word_list = text_box.get("1.0", "end-1c").splitlines()
    if not word_list:
        messagebox.showerror("Error", "Word list is empty.")
        return

    # Ask the user for the input file path
    docx_filename = filedialog.askopenfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
    if not docx_filename:
        return

    # Ask the user for the output file path
    output_filename = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
    if not output_filename:
        return

    # Append "_processed" to the filename
    output_filename = output_filename.replace(".docx", "_processed.docx")

    # Process the Word document and check if it was successful
    if process_word_document(docx_filename, output_filename, word_list):
        messagebox.showinfo("Success", f"Processing completed successfully. Result saved as '{output_filename}'")
    else:
        messagebox.showerror("Error", "Processing failed.")

# Function to update word list
def update_word_list():
    global saved_word_list
    saved_word_list = text_box.get("1.0", "end-1c").splitlines()

# Function to add a word or phrase to the saved word list
def add_word_or_phrase():
    global saved_word_list
    new_entry = simpledialog.askstring("Add Word/Phrase", "Enter a word or phrase to add:")
    if new_entry:
        saved_word_list.append(new_entry)
        update_text_box(saved_word_list)

# Function to remove a word or phrase from the saved word list
def remove_word_or_phrase():
    global saved_word_list
    entry_to_remove = simpledialog.askstring("Remove Word/Phrase", "Enter a word or phrase to remove:")
    if entry_to_remove in saved_word_list:
        saved_word_list.remove(entry_to_remove)
        update_text_box(saved_word_list)

# Function to update the text box with the word list
def update_text_box(word_list):
    text_box.delete("1.0", tk.END)
    for entry in word_list:
        text_box.insert(tk.END, entry + "\n")

def add_double_brackets_to_words_in_text(input_text, word_list):
    result = input_text

    for entry in word_list:
        # Check if the entry is a whole word or phrase (case-insensitive)
        if re.search(rf'\b{re.escape(entry)}\b', result, re.IGNORECASE):
            result = re.sub(rf'\b({re.escape(entry)})\b', r'[[\1]]', result, flags=re.IGNORECASE)

    return result

def process_word_document(input_filename, result_filename, word_list):
    try:
        # Open the original document in read-only mode
        doc = docx.Document(input_filename)

        # Create a new document for processing
        result_doc = docx.Document()

        # Process each paragraph from the original document
        for paragraph in doc.paragraphs:
            processed_paragraph = add_double_brackets_to_words_in_text(paragraph.text.strip(), word_list)
            result_doc.add_paragraph(processed_paragraph)

        # Save the result to the specified Word document
        result_doc.save(result_filename)
        return True

    except Exception as e:
        logging.error(f"Error: Unable to read or save the Word document. Error: {str(e)}")
        return False

# Create the main application window
app = tk.Tk()
app.title("Word Document Processor")
app.configure(bg="#808080")  # Set the background color to gray

# Create and configure widgets
label = tk.Label(app, text="Word List:", fg="white", bg="#808080", font=("Arial", 14))
text_box = tk.Text(app, height=10, width=40, font=("Arial", 12))
update_word_list_button = tk.Button(app, text="Update Word List", fg="white", bg="#007AFF", font=("Arial", 14), command=update_word_list)
process_button = tk.Button(app, text="Process", fg="white", bg="#007AFF", font=("Arial", 14), command=process_document)
add_word_button = tk.Button(app, text="Add Word/Phrase", fg="white", bg="#007AFF", font=("Arial", 14), command=add_word_or_phrase)
remove_word_button = tk.Button(app, text="Remove Word/Phrase", fg="white", bg="#007AFF", font=("Arial", 14), command=remove_word_or_phrase)

# Place widgets on the window
label.pack(pady=20)
text_box.pack(padx=20, pady=10)
update_word_list_button.pack(pady=10)
process_button.pack(pady=10)
add_word_button.pack(pady=10)
remove_word_button.pack(pady=10)

# Start the GUI application
app.mainloop()
