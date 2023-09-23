import docx
import os
import logging
import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
import re
import shutil
import glob

# Initialize the saved word list as empty
saved_word_list = []

# Set the initial working directory
initial_working_directory = r'C:\Users\HamburgerC\Desktop\Bracketing code'

# Variable to track whether processing should occur
processing_enabled = False

# Output folder for processed files
output_base_folder = r'C:\Users\HamburgerC\Desktop\Bracketing code\Processed'

# Output folder for processed documents
output_documents_folder = os.path.join(output_base_folder, 'Processed Documents')

# Output folder for processed folders
output_folders_folder = os.path.join(output_base_folder, 'Processed Folders')

# Create output directories if they don't exist
os.makedirs(output_documents_folder, exist_ok=True)
os.makedirs(output_folders_folder, exist_ok=True)

# Function to select a folder
def select_folder():
    folder_path = filedialog.askdirectory(initialdir=initial_working_directory)
    if folder_path:
        process_folder(folder_path)

# Function to select a Word document
def select_document():
    doc_path = filedialog.askopenfilename(initialdir=initial_working_directory)
    if doc_path:
        process_document(doc_path)

# Function to add double brackets to words in text (with a maximum of two brackets)
def add_double_brackets_to_words_in_text(input_text, word_list):
    result_text = input_text

    for word in word_list:
        # Check if the word already contains double brackets
        if f"[[{word}]]" not in result_text:
            # Use a case-sensitive replacement
            result_text = result_text.replace(word, f"[[{word}]]")
        else:
            # If double brackets are already present, replace any occurrences of triple brackets with double brackets
            result_text = result_text.replace(f"[[[{word}]]]", f"[[{word}]]")

    return result_text

# Function to extract text from a Word document
def extract_text_from_docx(docx_filename):
    try:
        doc = docx.Document(docx_filename)
        text = '\n'.join([paragraph.text.strip() for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        logging.error(f"Error: Unable to read the Word document. Error: {str(e)}")
        return ""

# Function to clean up extra brackets beside the two that should be there
def clean_up_extra_brackets(input_text):
    cleaned_text = input_text

    # Use a regular expression to remove extra brackets beside the two that should be there
    cleaned_text = re.sub(r'\[\[+([^\[\]]+)\]\]+', r'[[\1]]', cleaned_text)

    return cleaned_text

# Function to process a Word document
def process_word_document(input_filename, word_list, result_filename, num_iterations=3):
    try:
        # Load the Word document
        doc_text = extract_text_from_docx(input_filename)

        # Initialize an empty list to store the processed paragraphs
        processed_paragraphs = []

        for _ in range(num_iterations):
            # Process the document text
            processed_text = add_double_brackets_to_words_in_text(doc_text, word_list)
            processed_text = clean_up_extra_brackets(processed_text)
            processed_paragraphs.append(processed_text)

        result_text = '\n'.join(processed_paragraphs)

        # Create a sanitized output file name by replacing problematic characters
        output_filename = re.sub(r'[\/:*?"<>|]', '_', os.path.basename(result_filename))
        full_output_path = os.path.join(result_filename, output_filename)

        # Check if the file already exists and prompt for overwrite
        if os.path.exists(full_output_path):
            confirm = messagebox.askyesno("File Exists", f"The file '{full_output_path}' already exists. Do you want to overwrite it?")
            if not confirm:
                logging.error("Processing canceled.")
                return False

        # Create a new Word document to store the result
        result_doc = docx.Document()
        result_doc.add_paragraph(result_text)

        # Save the result to the specified Word document
        result_doc.save(full_output_path)

        return True
    except Exception as e:
        logging.error(f"Error: An error occurred while processing the document. Error: {str(e)}")
        return False

# Function to check for and limit the number of brackets around words (maximum of two brackets)
def limit_brackets(word_list):
    for i in range(len(word_list)):
        word = word_list[i]
        # Use a while loop to remove extra brackets until only a maximum of two brackets remain
        while "[[" in word:
            word = word.replace("[[", "[")
        while "]]" in word:
            word = word.replace("]]", "]")
        word_list[i] = word

# Function to process an individual Word document
def process_document(doc_path):
    try:
        global processing_enabled  # Access the global variable

        if processing_enabled:  # Check if processing is enabled
            # Create a sanitized output file name by replacing problematic characters
            output_filename = re.sub(r'[\/:*?"<>|]', '_', doc_path)
            output_filename = os.path.join(output_documents_folder, output_filename + "_Processed.docx")

            # Process the Word document and display a message based on success or failure
            if process_word_document(doc_path, saved_word_list, output_filename):
                messagebox.showinfo("Success", f"Processing completed successfully. Result saved as '{output_filename}'")
            else:
                messagebox.showerror("Error", "Processing failed.")
    except Exception as e:
        logging.error(f"Error: An error occurred while processing the document. Error: {str(e)}")

# Function to process a file and double bracket its content
def process_file(file_path):
    try:
        global processing_enabled  # Access the global variable

        if processing_enabled:  # Check if processing is enabled
            # Read the file content
            with open(file_path, 'r', encoding='utf-8') as file_content:
                file_text = file_content.read()

            # Create a sanitized output file name by replacing problematic characters
            output_filename = re.sub(r'[\/:*?"<>|]', '_', os.path.basename(file_path))
            output_filename = os.path.join(output_documents_folder, output_filename + ".txt")

            # Double bracket the content and save it to the specified file
            result_text = add_double_brackets_to_words_in_text(file_text, saved_word_list)

            # Save the result to the specified file
            with open(output_filename, 'w', encoding='utf-8') as result_file:
                result_file.write(result_text)

            messagebox.showinfo("Success", f"Processing completed successfully. Result saved as '{output_filename}'")
    except Exception as e:
        logging.error(f"Error: An error occurred while processing the file. Error: {str(e)}")

# Function to update word list
def update_word_list(text_box):
    global saved_word_list, processing_enabled
    saved_word_list = text_box.get("1.0", "end-1c").splitlines()
    # Check and limit the number of brackets around words
    limit_brackets(saved_word_list)
    
    # Enable processing if there are words in the list
    if saved_word_list:
        processing_enabled = True
        select_doc_button.config(state=tk.NORMAL, fg='green')
        select_folder_button.config(state=tk.NORMAL, fg='green')
    else:
        processing_enabled = False
        select_doc_button.config(state=tk.DISABLED)
        select_folder_button.config(state=tk.DISABLED)
        messagebox.showwarning("Warning", "Word list is empty. Processing is disabled.")

# Function to update the text box with the word list
def update_text_box(word_list, text_box):
    text_box.delete("1.0", tk.END)
    for word in word_list:
        text_box.insert(tk.END, word + "\n")

# Function to save the word list to a text file
def save_word_list(text_box):
    global saved_word_list
    file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt")], initialdir=initial_working_directory)
    if file_path:
        with open(file_path, "w") as file:
            for word in saved_word_list:
                file.write(word + "\n")

# Function to load the word list from a text file
def load_word_list(text_box):
    global saved_word_list
    file_path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt")], initialdir=initial_working_directory)
    if file_path:
        with open(file_path, "r") as file:
            saved_word_list = [line.strip() for line in file]
            # Check and limit the number of brackets around words
            limit_brackets(saved_word_list)
            update_text_box(saved_word_list, text_box)  # Update the text box with the loaded word list

# Function to change the working directory to the specified path
def change_working_directory(desired_path):
    os.chdir(desired_path)

# Function to duplicate a folder and its contents
def duplicate_folder(src_folder, dest_folder):
    try:
        # Create the destination folder if it doesn't exist
        if not os.path.exists(dest_folder):
            os.makedirs(dest_folder)

        # Iterate over all files and subdirectories in the source folder
        for item in os.listdir(src_folder):
            src_item = os.path.join(src_folder, item)
            dest_item = os.path.join(dest_folder, item)

            # If the item is a file, copy it to the destination folder
            if os.path.isfile(src_item):
                shutil.copy2(src_item, dest_item)

            # If the item is a directory, recursively duplicate it
            elif os.path.isdir(src_item):
                duplicate_folder(src_item, dest_item)

    except Exception as e:
        logging.error(f"Error: An error occurred while duplicating the folder. Error: {str(e)}")

# Function to process all Markdown files within a folder and its subfolders
def process_folder(folder_path):
    try:
        # Create a destination folder for processed documents
        processed_folder = os.path.join(folder_path, "Processed_Documents")
        if not os.path.exists(processed_folder):
            os.makedirs(processed_folder)

        # Search for Markdown files within the folder and its subfolders
        md_files = glob.glob(os.path.join(folder_path, '**/*.md'), recursive=True)

        # Process each Markdown file and create a corresponding Word document
        for md_file in md_files:
            process_markdown_file(md_file, saved_word_list)

        messagebox.showinfo("Success", f"Processing completed successfully. Processed files saved in '{processed_folder}'")

    except Exception as e:
        logging.error(f"Error: An error occurred while processing the folder. Error: {str(e)}")

# Function to process a Markdown file
def process_markdown_file(md_file, word_list):
    try:
        # Determine the relative path within the folder
        relative_path = os.path.relpath(md_file, initial_working_directory)

        # Create a sanitized output folder path by replacing problematic characters
        sanitized_relative_path = re.sub(r'[\/:*?"<>|]', '_', relative_path)
        output_folder_path = os.path.join(output_folders_folder, os.path.splitext(sanitized_relative_path)[0])

        # Create the output folder if it doesn't exist
        os.makedirs(output_folder_path, exist_ok=True)

        # Create the corresponding output file path in the processed folder with a .docx extension
        output_file_name = os.path.splitext(os.path.basename(md_file))[0] + ".docx"
        output_file_path = os.path.join(output_folder_path, output_file_name)

        # Read the Markdown file content
        with open(md_file, 'r', encoding='utf-8') as md_file_content:
            md_content = md_file_content.read()

        # Process the Markdown content
        processed_md_content = add_double_brackets_to_words_in_text(md_content, word_list)

        # Create a Word document and add the processed content
        result_doc = docx.Document()
        result_doc.add_paragraph(processed_md_content)

        # Save the result to the corresponding Word document
        result_doc.save(output_file_path)
    except Exception as e:
        logging.error(f"Error: An error occurred while processing the Markdown file. Error: {str(e)}")

# Function to start the GUI application
def start_gui():
    global select_doc_button, select_folder_button  # Declare the buttons as global variables
    # Create the main application window
    app = tk.Tk()
    app.title("Markdown to Word Document Processor")

    # Create and configure widgets
    text_box = tk.Text(app, height=10, width=40, font=("Arial", 12))
    text_box.pack(padx=20, pady=10)

    button_frame = tk.Frame(app)
    button_frame.pack(pady=20)

    # Buttons to select a Word document and a folder
    select_doc_button = tk.Button(button_frame, text="Select Document", font=("Arial", 14), command=select_document)
    select_doc_button.grid(row=0, column=0, padx=5)
    
    select_folder_button = tk.Button(button_frame, text="Select Folder", font=("Arial", 14), command=select_folder)
    select_folder_button.grid(row=0, column=1, padx=5)

    update_word_list_button = tk.Button(button_frame, text="Update Word List", font=("Arial", 14), command=lambda: update_word_list(text_box))
    update_word_list_button.grid(row=1, column=0, columnspan=2, pady=10)
    
    # Set the initial state of document and folder buttons
    select_doc_button.config(state=tk.DISABLED)
    select_folder_button.config(state=tk.DISABLED)

    save_word_list_button = tk.Button(button_frame, text="Save Word List", font=("Arial", 14), command=lambda: save_word_list(text_box))
    save_word_list_button.grid(row=2, column=0, padx=5)
    
    load_word_list_button = tk.Button(button_frame, text="Load Word List", font=("Arial", 14), command=lambda: load_word_list(text_box))
    load_word_list_button.grid(row=2, column=1, padx=5)

    try:
        app.mainloop()
    except Exception as e:
        print(f"An error occurred: {str(e)}")

# Start the GUI application
start_gui()
