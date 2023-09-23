import docx
import os
import logging

# Configure logging without creating a log file
logging.basicConfig(level=logging.ERROR)

# Initialize the last used word list as an empty list
last_used_word_list = []

def add_double_brackets_to_words_in_text(input_text, word_list):
    words = input_text.split()
    result = []

    for word in words:
        # Check if the word matches any word in the provided list (case-insensitive)
        if any(word.lower() == item.lower() for item in word_list):
            result.append(f"[[{word}]]")  # Double brackets
        else:
            result.append(word)

    result_text = ' '.join(result)
    return result_text

def extract_text_from_docx(docx_filename):
    try:
        doc = docx.Document(docx_filename)
        text = '\n'.join([paragraph.text.strip() for paragraph in doc.paragraphs])
        return text
    except FileNotFoundError as e:
        logging.error(f"Error: The file '{docx_filename}' was not found. Error: {str(e)}")
        return ""

def process_word_document(docx_filename, word_list, result_filename):
    try:
        # Load the Word document
        doc_text = extract_text_from_docx(docx_filename)
    except FileNotFoundError as e:
        logging.error(f"Error: The file '{docx_filename}' was not found. Error: {str(e)}")
        return False

    # Split the document text into paragraphs
    paragraphs = doc_text.split('\n')

    # Initialize an empty list to store the processed paragraphs
    processed_paragraphs = []

    # Process each paragraph
    for paragraph in paragraphs:
        processed_paragraph = add_double_brackets_to_words_in_text(paragraph, word_list)
        processed_paragraphs.append(processed_paragraph)

    # Combine the processed paragraphs into a single text
    result_text = '\n'.join(processed_paragraphs)

    # Create a new Word document to store the result
    result_doc = docx.Document()
    result_doc.add_paragraph(result_text)

    # Save the result to the specified Word document
    result_doc.save(result_filename)
    return True

def print_menu():
    print("1. Enter a new word list")
    print("2. Use the last word list")
    print("3. Quit")

if __name__ == "__main__":
    while True:
        print_menu()
        choice = input("Enter your choice: ")

        if choice == "1":
            # Prompt the user to enter the word list
            word_list_text = input("Enter the word list (one word per line, end with an empty line):\n")

            # Split the word list text into individual words
            last_used_word_list = word_list_text.strip().split('\n')
            print("Word list updated.")

        elif choice == "2":
            if last_used_word_list:
                # Prompt the user for the output file path
                result_filename = input("Enter the path for the output Word document: ")

                # Process the Word document and check if it was successful
                if process_word_document(word_list_path, last_used_word_list, result_filename):
                    print(f"Processing completed successfully. Result saved as '{result_filename}'")
                else:
                    print("Processing failed.")
            else:
                print("No word list available. Please enter a new word list (Option 1).")

        elif choice == "3":
            break

        else:
            print("Invalid choice. Please select a valid option (1, 2, or 3).")
