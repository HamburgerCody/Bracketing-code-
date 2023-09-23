import docx
import os
import logging
import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox

# Configure logging without creating a log file
logging.basicConfig(level=logging.ERROR)

# Initialize the last used word list as empty
last_used_word_list = []

def add_double_brackets_to_words_in_text(input_text, word_list):
    words = input_text.split()
    result = []

    for word in words:
        # Check if the word matches any word in the provided list (case-insensitive)
        if any(word.lower() == item.lower() for item in word_list):
            result.append(f"[[{word}]]")  # Double brackets
        else:
            result.append(word)

    result_text = ' '.join(result)
    return result_text

def extract_text_from_docx(docx_filename):
    try:
        doc = docx.Document(docx_filename)
        text = '\n'.join([paragraph.text.strip() for paragraph in doc.paragraphs])
        return text
    except FileNotFoundError as e:
        logging.error(f"Error: The file '{docx_filename}' was not found. Error: {str(e)}")
        return ""

def process_word_document(docx_filename, word_list, result_filename):
    try:
        # Load the Word document
        doc_text = extract_text_from_docx(docx_filename)
    except FileNotFoundError as e:
        logging.error(f"Error: The file '{docx_filename}' was not found. Error: {str(e)}")
        return False

    # Split the document text into paragraphs
    paragraphs = doc_text.split('\n')

    # Initialize an empty list to store the processed paragraphs
    processed_paragraphs = []

    # Process each paragraph
    for paragraph in paragraphs:
        processed_paragraph = add_double_brackets_to_words_in_text(paragraph, word_list)
        processed_paragraphs.append(processed_paragraph)

    # Combine the processed paragraphs into a single text
    result_text = '\n'.join(processed_paragraphs)

    # Determine the full output path
    full_output_path = os.path.join(result_filename)

    # Check if the file already exists and prompt for overwrite
    if os.path.exists(full_output_path):
        confirm = messagebox.askyesno("File Exists", f"The file '{full_output_path}' already exists. Do you want to overwrite it?")
        if not confirm:
            print("Processing canceled.")
            return False

    # Create a new Word document to store the result
    result_doc = docx.Document()
    result_doc.add_paragraph(result_text)

    # Save the result to the specified Word document
    result_doc.save(full_output_path)
    return True

# Create the main application window
app = tk.Tk()
app.title("Word Document Processor")

# Create and configure widgets
label = tk.Label(app, text="Word List:")
text_box = tk.Text(app, height=10, width=40)
process_button = tk.Button(app, text="Process", command=process_document)
update_word_list_button = tk.Button(app, text="Update Word List", command=update_word_list)

# Place widgets on the window
label.pack(padx=10, pady=10)
text_box.pack(padx=10, pady=10)
update_word_list_button.pack(padx=10, pady=10)
process_button.pack(padx=10, pady=10)

# Function to process the document
def process_document():
    global last_used_word_list
    word_list = text_box.get("1.0", "end-1c").splitlines()
    if not word_list:
        messagebox.showerror("Error", "Word list is empty.")
        return

    output_filename = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
    if not output_filename:
        return

    # Process the Word document and check if it was successful
    if process_word_document(output_filename, word_list, output_filename):
        messagebox.showinfo("Success", f"Processing completed successfully. Result saved as '{output_filename}'")
    else:
        messagebox.showerror("Error", "Processing failed.")

# Function to update word list
def update_word_list():
    global last_used_word_list
    last_used_word_list = text_box.get("1.0", "end-1c").splitlines()

# Start the GUI application
app.mainloop()
