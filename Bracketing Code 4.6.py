import docx
import os
import logging
import tkinter as tk
from tkinter import filedialog

# Configure logging without creating a log file
logging.basicConfig(level=logging.ERROR)

# Initialize the last used word list and output path as empty
last_used_word_list = []
last_output_path = ""

def add_double_brackets_to_words_in_text(input_text, word_list):
    words = input_text.split()
    result = []

    for word in words:
        # Check if the word matches any word in the provided list (case-insensitive)
        if any(word.lower() == item.lower() for item in word_list):
            result.append(f"[[{word}]]")  # Double brackets
        else:
            result.append(word)

    result_text = ' '.join(result)
    return result_text

def extract_text_from_docx(docx_filename):
    try:
        doc = docx.Document(docx_filename)
        text = '\n'.join([paragraph.text.strip() for paragraph in doc.paragraphs])
        return text
    except FileNotFoundError as e:
        logging.error(f"Error: The file '{docx_filename}' was not found. Error: {str(e)}")
        return ""

def process_word_document(docx_filename, word_list, result_filename):
    try:
        # Load the Word document
        doc_text = extract_text_from_docx(docx_filename)
    except FileNotFoundError as e:
        logging.error(f"Error: The file '{docx_filename}' was not found. Error: {str(e)}")
        return False

    # Split the document text into paragraphs
    paragraphs = doc_text.split('\n')

    # Initialize an empty list to store the processed paragraphs
    processed_paragraphs = []

    # Process each paragraph
    for paragraph in paragraphs:
        processed_paragraph = add_double_brackets_to_words_in_text(paragraph, word_list)
        processed_paragraphs.append(processed_paragraph)

    # Combine the processed paragraphs into a single text
    result_text = '\n'.join(processed_paragraphs)

    # Create a new Word document to store the result
    result_doc = docx.Document()
    result_doc.add_paragraph(result_text)

    # Save the result to the specified Word document
    result_doc.save(result_filename)
    return True

def set_output_path():
    global last_output_path
    last_output_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])

def process_using_last_settings():
    global last_output_path
    if last_used_word_list and last_output_path:
        # Process the Word document using the last word list and output path
        if process_word_document(last_output_path, last_used_word_list, last_output_path):
            print(f"Processing completed successfully. Result saved as '{last_output_path}'")
        else:
            print("Processing failed.")
    else:
        print("No word list or output path set. Please set them.")

def update_word_list(word_list_text):
    global last_used_word_list
    last_used_word_list = word_list_text.split()

# Create the main application window
app = tk.Tk()
app.title("Word Document Processor")

# Create and configure widgets
label = tk.Label(app, text="Word List:")
text_box = tk.Text(app, height=10, width=40)
process_button = tk.Button(app, text="Process Using Last Settings", command=process_using_last_settings)
set_output_button = tk.Button(app, text="Set Output Path", command=set_output_path)
update_word_list_button = tk.Button(app, text="Update Word List", command=lambda: update_word_list(text_box.get("1.0", "end-1c")))

# Place widgets on the window
label.grid(row=0, column=0, padx=10, pady=10)
text_box.grid(row=0, column=1, padx=10, pady=10)
update_word_list_button.grid(row=1, column=0, padx=10, pady=10)
set_output_button.grid(row=1, column=1, padx=10, pady=10)
process_button.grid(row=2, column=0, columnspan=2, padx=10, pady=10)

# Start the GUI application
app.mainloop()
