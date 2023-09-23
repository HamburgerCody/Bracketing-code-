import docx
import os
import logging

# Define constants for file paths
DOCX_FILENAME = r"C:\Users\HamburgerC\Desktop\Bracketing code\Unbracketed.docx"
WORD_LIST_PATH = r"C:\Users\HamburgerC\Desktop\Bracketing code\Word List\Finley Thornwood\New Text Document.txt"
RESULT_FILENAME = r"C:\Users\HamburgerC\Desktop\Bracketing code\Processed_Bracketed.docx"

# Configure logging without creating a log file
logging.basicConfig(level=logging.ERROR)

def add_double_brackets_to_words_in_text(input_text, word_list):
    words = input_text.split()
    result = []

    for word in words:
        # Check if the word matches any word in the provided list (case-insensitive)
        if any(word.lower() == item.lower() for item in word_list):
            result.append(f"[[{word}]]")  # Double brackets
        else:
            result.append(word)

    result_text = ' '.join(result)
    return result_text

def read_word_list_from_txt(txt_filename):
    try:
        with open(txt_filename, 'r', encoding='utf-8') as file:
            word_list = [line.strip() for line in file.readlines()]
        return word_list
    except FileNotFoundError:
        logging.error(f"Error: The file '{txt_filename}' was not found.")
        return []

def process_word_document(docx_filename, word_list, result_filename):
    try:
        # Load the Word document
        doc = docx.Document(docx_filename)
    except FileNotFoundError:
        logging.error(f"Error: The file '{docx_filename}' was not found.")
        return False

    # Initialize an empty list to store the paragraphs
    paragraphs = []

    # Iterate through the paragraphs in the document
    for paragraph in doc.paragraphs:
        # Process each paragraph's text and add it to the list
        processed_paragraph = add_double_brackets_to_words_in_text(paragraph.text, word_list)
        paragraphs.append(processed_paragraph)

    # Create a new Word document to store the result
    result_doc = docx.Document()

    # Add processed paragraphs to the result document
    for paragraph in paragraphs:
        result_doc.add_paragraph(paragraph)

    # Check if the result file already exists
    if os.path.isfile(result_filename):
        overwrite = input("The result file already exists. Do you want to overwrite it? (yes/no): ")
        if overwrite.lower() != 'yes':
            print("Operation canceled.")
            return False

    # Save the result to a new Word document
    result_doc.save(result_filename)
    return True

if __name__ == "__main__":
    # Read the word list from the text file
    word_list = read_word_list_from_txt(WORD_LIST_PATH)

    if not word_list:
        print("Word list is empty or could not be read.")
    else:
        # Process the Word document and check if it was successful
        if process_word_document(DOCX_FILENAME, word_list, RESULT_FILENAME):
            print(f"Processing completed successfully. Result saved as '{RESULT_FILENAME}'")
        else:
            print("Processing failed.")
