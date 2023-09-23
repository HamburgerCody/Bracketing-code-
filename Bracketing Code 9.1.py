import docx
import os
import logging
import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
from tkinter import simpledialog  # For user input

# Configure logging without creating a log file
logging.basicConfig(level=logging.ERROR)

# Initialize the saved word list as empty
saved_word_list = []

# Function to add double brackets to words in text
def add_double_brackets_to_words_in_text(input_text, word_list):
    words = input_text.split()
    result = []

    for word in words:
        # Check if the word matches any word in the provided list (case-insensitive)
        if any(word.lower() == item.lower() for item in word_list):
            result.append(f"[[{word}]]")  # Double brackets
        else:
            result.append(word)

    result_text = ' '.join(result)
    return result_text

# Function to extract text from a Word document
def extract_text_from_docx(docx_filename):
    try:
        doc = docx.Document(docx_filename)
        text = '\n'.join([paragraph.text.strip() for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        logging.error(f"Error: Unable to read the Word document. Error: {str(e)}")
        return ""

# Function to process a Word document
def process_word_document(input_filename, word_list, result_filename, num_iterations=3):
    try:
        # Load the Word document
        doc_text = extract_text_from_docx(input_filename)

        # Split the document text into paragraphs
        paragraphs = doc_text.split('\n')

        # Initialize an empty list to store the processed paragraphs
        processed_paragraphs = []

        for _ in range(num_iterations):
            # Process each paragraph
            for paragraph in paragraphs:
                processed_paragraph = add_double_brackets_to_words_in_text(paragraph, word_list)
                processed_paragraphs.append(processed_paragraph)

            # Combine the processed paragraphs into a single text
            paragraphs = processed_paragraphs
            processed_paragraphs = []

        result_text = '\n'.join(paragraphs)

        # Save the result to the specified Word document without adding "_processed"
        full_output_path = os.path.abspath(result_filename)

        # Check if the file already exists and prompt for overwrite
        if os.path.exists(full_output_path):
            confirm = messagebox.askyesno("File Exists", f"The file '{full_output_path}' already exists. Do you want to overwrite it?")
            if not confirm:
                logging.error("Processing canceled.")
                return False

        # Create a new Word document to store the result
        result_doc = docx.Document()
        result_doc.add_paragraph(result_text)

        # Save the result to the specified Word document
        result_doc.save(full_output_path)
        return True
    except Exception as e:
        logging.error(f"Error: An error occurred while processing the document. Error: {str(e)}")
        return False

# Function to process the document
def process_document():
    global saved_word_list
    word_list = text_box.get("1.0", "end-1c").splitlines()
    if not word_list:
        messagebox.showerror("Error", "Word list is empty.")
        return

    input_filename = filedialog.askopenfilename(filetypes=[("Word Document", "*.docx")])
    if not input_filename:
        return

    output_filename = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
    if not output_filename:
        return

    num_iterations = 3  # You can change this to control the number of processing iterations

    # Process the Word document and display a message based on success or failure
    if process_word_document(input_filename, word_list, output_filename, num_iterations):
        messagebox.showinfo("Success", f"Processing completed successfully. Result saved as '{output_filename}'")
    else:
        messagebox.showerror("Error", "Processing failed.")

# Function to update word list
def update_word_list():
    global saved_word_list
    saved_word_list = text_box.get("1.0", "end-1c").splitlines()

# Function to add a word to the saved word list
def add_word():
    global saved_word_list
    new_word = simpledialog.askstring("Add Word", "Enter a word to add:")
    if new_word:
        saved_word_list.append(new_word)
        update_text_box(saved_word_list)

# Function to remove a word from the saved word list
def remove_word():
    global saved_word_list
    word_to_remove = simpledialog.askstring("Remove Word", "Enter a word to remove:")
    if word_to_remove in saved_word_list:
        saved_word_list.remove(word_to_remove)
        update_text_box(saved_word_list)

# Function to update the text box with the word list
def update_text_box(word_list):
    text_box.delete("1.0", tk.END)
    for word in word_list:
        text_box.insert(tk.END, word + "\n")

# Create the main application window
app = tk.Tk()
app.title("Word Document Processor")
app.configure(bg="#808080")  # Set the background color to gray

# Create and configure widgets
label = tk.Label(app, text="Word List:", fg="white", bg="#808080", font=("Arial", 14))
text_box = tk.Text(app, height=10, width=40, font=("Arial", 12))
update_word_list_button = tk.Button(app, text="Update Word List", fg="white", bg="#007AFF", font=("Arial", 14), command=update_word_list)
process_button = tk.Button(app, text="Process", fg="white", bg="#007AFF", font=("Arial", 14), command=process_document)
add_word_button = tk.Button(app, text="Add Word", fg="white", bg="#007AFF", font=("Arial", 14), command=add_word)
remove_word_button = tk.Button(app, text="Remove Word", fg="white", bg="#007AFF", font=("Arial", 14), command=remove_word)

# Place widgets on the window
label.pack(pady=20)
text_box.pack(padx=20, pady=10)
update_word_list_button.pack(pady=10)
process_button.pack(pady=10)
add_word_button.pack(pady=10)
remove_word_button.pack(pady=10)

# Start the GUI application
app.mainloop()
