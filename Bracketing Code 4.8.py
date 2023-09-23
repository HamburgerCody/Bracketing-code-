import docx
import os
import logging
import tkinter as tk
from tkinter import filedialog

# Configure logging without creating a log file
logging.basicConfig(level=logging.ERROR)

# Initialize the last used word list and output path as empty
last_used_word_list = []
output_directory = r'C:\Users\HamburgerC\Desktop\Bracketing code'

def add_double_brackets_to_words_in_text(input_text, word_list):
    words = input_text.split()
    result = []

    for word in words:
        # Check if the word matches any word in the provided list (case-insensitive)
        if any(word.lower() == item.lower() for item in word_list):
            result.append(f"[[{word}]]")  # Double brackets
        else:
            result.append(word)

    result_text = ' '.join(result)
    return result_text

def extract_text_from_docx(docx_filename):
    try:
        doc = docx.Document(docx_filename)
        text = '\n'.join([paragraph.text.strip() for paragraph in doc.paragraphs])
        return text
    except FileNotFoundError as e:
        logging.error(f"Error: The file '{docx_filename}' was not found. Error: {str(e)}")
        return ""

def process_word_document(docx_filename, word_list, result_filename):
    try:
        # Load the Word document
        doc_text = extract_text_from_docx(docx_filename)
    except FileNotFoundError as e:
        logging.error(f"Error: The file '{docx_filename}' was not found. Error: {str(e)}")
        return False

    # Split the document text into paragraphs
    paragraphs = doc_text.split('\n')

    # Initialize an empty list to store the processed paragraphs
    processed_paragraphs = []

    # Process each paragraph
    for paragraph in paragraphs:
        processed_paragraph = add_double_brackets_to_words_in_text(paragraph, word_list)
        processed_paragraphs.append(processed_paragraph)

    # Combine the processed paragraphs into a single text
    result_text = '\n'.join(processed_paragraphs)

    # Determine the full output path
    full_output_path = os.path.join(output_directory, result_filename)

    # Check if the file already exists and prompt for overwrite
    if os.path.exists(full_output_path):
        overwrite = input(f"The file '{full_output_path}' already exists. Do you want to overwrite it? (y/n): ")
        if overwrite.lower() != 'y':
            print("Processing canceled.")
            return False

    # Create a new Word document to store the result
    result_doc = docx.Document()
    result_doc.add_paragraph(result_text)

    # Save the result to the specified Word document
    result_doc.save(full_output_path)
    return True

# Create the main application window
app = tk.Tk()
app.title("Word Document Processor")

# Create and configure widgets
label = tk.Label(app, text="Word List:")
text_box = tk.Text(app, height=10, width=40)
process_button = tk.Button(app, text="Process Using Last Settings", command=process_using_last_settings)
update_word_list_button = tk.Button(app, text="Update Word List", command=lambda: update_word_list(text_box.get("1.0", "end-1c")))
quit_button = tk.Button(app, text="Quit", command=app.quit)  # Added Quit button

# Place widgets on the window
label.grid(row=0, column=0, padx=10, pady=10)
text_box.grid(row=0, column=1, padx=10, pady=10)
update_word_list_button.grid(row=1, column=0, padx=10, pady=10)
process_button.grid(row=1, column=1, padx=10, pady=10)
quit_button.grid(row=2, column=0, columnspan=2, padx=10, pady=10)  # Added Quit button

# Function to process using last settings
def process_using_last_settings():
    global last_used_word_list
    if last_used_word_list:
        # Prompt the user for the output file name
        result_filename = input("Enter the output file name (e.g., 'Output.docx'): ")

        # Process the Word document and check if it was successful
        if process_word_document(result_filename, last_used_word_list, result_filename):
            print(f"Processing completed successfully. Result saved as '{result_filename}'")
        else:
            print("Processing failed.")
    else:
        print("No word list available. Please enter a new word list (Option 1).")

# Function to update word list
def update_word_list(word_list_text):
    global last_used_word_list
    last_used_word_list = word_list_text.split()

# Start the GUI application
app.mainloop()
