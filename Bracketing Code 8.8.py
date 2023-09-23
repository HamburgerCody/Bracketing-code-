import docx
import os
import logging
import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
from tkinter import simpledialog  # For user input

# Configure logging without creating a log file
logging.basicConfig(level=logging.ERROR)

# Initialize the saved word list as empty
saved_word_list = []

def add_double_brackets_to_words_in_text(input_text, word_list):
    words = input_text.split()
    result = []

    for word in words:
        # Check if the word matches any word in the provided list (case-insensitive)
        if any(word.lower() == item.lower() for item in word_list):
            result.append(f"[[{word}]]")  # Double brackets
        else:
            result.append(word)

    result_text = ' '.join(result)
    return result_text

def extract_text_from_docx(docx_filename):
    try:
        doc = docx.Document(docx_filename)
        text = '\n'.join([paragraph.text.strip() for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        logging.error(f"Error: Unable to read the Word document. Error: {str(e)}")
        return ""

def process_word_document(input_filename, word_list, result_filename):
    try:
        # Load the Word document
        doc_text = extract_text_from_docx(input_filename)

        # Split the document text into paragraphs
        paragraphs = doc_text.split('\n')

        # Initialize an empty list to store the processed paragraphs
        processed_paragraphs = []

        # Number of processing iterations
        num_iterations = 3

        for _ in range(num_iterations):
            # Process each paragraph
            for paragraph in paragraphs:
                processed_paragraph = add_double_brackets_to_words_in_text(paragraph, word_list)
                processed_paragraphs.append(processed_paragraph)

            # Combine the processed paragraphs into a single text
            paragraphs = processed_paragraphs  # Update paragraphs for the next iteration
            processed_paragraphs = []  # Reset the processed paragraphs

        # Combine the final processed paragraphs into a single text
        result_text = '\n'.join(paragraphs)

        # Save the result to the specified Word document without adding "_processed"
        full_output_path = os.path.abspath(result_filename)

        # Check if the file already exists and prompt for overwrite
        if os.path.exists(full_output_path):
            confirm = messagebox.askyesno("File Exists", f"The file '{full_output_path}' already exists. Do you want to overwrite it?")
            if not confirm:
                logging.error("Processing canceled.")
                return False

        # Create a new Word document to store the result
        result_doc = docx.Document()
        result_doc.add_paragraph(result_text)

        # Save the result to the specified Word document
        result_doc.save(full_output_path)
        return True
    except Exception as e:
        logging.error(f"Error: An error occurred while processing the document. Error: {str(e)}")
        return False

# Create the main application window
app = tk.Tk()
app.title("Word Document Processor")
app.configure(bg="#808080")  # Set the background color to gray

# Rest of the code remains unchanged
