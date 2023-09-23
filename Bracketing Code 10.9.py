import docx
import os
import logging
import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
from tkinter import simpledialog  # For user input
from tkinter import ttk  # For themed buttons

# Configure logging without creating a log file
logging.basicConfig(level=logging.ERROR)

# Initialize the saved word list as empty
saved_word_list = []

# Function to add double brackets to words in text (with a maximum of two brackets)
def add_double_brackets_to_words_in_text(input_text, word_list):
    result_text = input_text
    
    for word in word_list:
        # Check if the word already contains double brackets
        if f"[[{word}]]" not in result_text:
            # Use a case-sensitive replacement
            result_text = result_text.replace(word, f"[[{word}]]")
        
        # Check if there are more than two brackets on the word and adjust if necessary
        result_text = result_text.replace(f"[[[{word}]]]", f"[[{word}]]")
    
    return result_text

# Function to extract text from a Word document
def extract_text_from_docx(docx_filename):
    try:
        doc = docx.Document(docx_filename)
        text = '\n'.join([paragraph.text.strip() for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        logging.error(f"Error: Unable to read the Word document. Error: {str(e)}")
        return ""

# Function to process a Word document
def process_word_document(input_filename, word_list, result_filename, num_iterations=3):
    try:
        # Load the Word document
        doc_text = extract_text_from_docx(input_filename)

        # Split the document text into paragraphs
        paragraphs = doc_text.split('\n')

        # Initialize an empty list to store the processed paragraphs
        processed_paragraphs = []

        for _ in range(num_iterations):
            # Process each paragraph
            for paragraph in paragraphs:
                processed_paragraph = add_double_brackets_to_words_in_text(paragraph, word_list)
                processed_paragraphs.append(processed_paragraph)

            # Combine the processed paragraphs into a single text
            paragraphs = processed_paragraphs
            processed_paragraphs = []

        result_text = '\n'.join(paragraphs)

        # Save the result to the specified Word document without adding "_processed"
        full_output_path = os.path.abspath(result_filename)

        # Check if the file already exists and prompt for overwrite
        if os.path.exists(full_output_path):
            confirm = messagebox.askyesno("File Exists", f"The file '{full_output_path}' already exists. Do you want to overwrite it?")
            if not confirm:
                logging.error("Processing canceled.")
                return False

        # Create a new Word document to store the result
        result_doc = docx.Document()
        result_doc.add_paragraph(result_text)

        # Save the result to the specified Word document
        result_doc.save(full_output_path)
        return True
    except Exception as e:
        logging.error(f"Error: An error occurred while processing the document. Error: {str(e)}")
        return False

# Function to process the document
def process_document():
    global saved_word_list
    word_list = text_box.get("1.0", "end-1c").splitlines()
    if not word_list:
        messagebox.showerror("Error", "Word list is empty.")
        return

    input_filename = filedialog.askopenfilename(filetypes=[("Word Document", "*.docx")])
    if not input_filename:
        return

    output_filename = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
    if not output_filename:
        return

    num_iterations = 3  # You can change this to control the number of processing iterations

    # Process the Word document and display a message based on success or failure
    if process_word_document(input_filename, word_list, output_filename, num_iterations):
        messagebox.showinfo("Success", f"Processing completed successfully. Result saved as '{output_filename}'")
    else:
        messagebox.showerror("Error", "Processing failed.")

# Function to update word list
def update_word_list():
    global saved_word_list
    saved_word_list = text_box.get("1.0", "end-1c").splitlines()

# Function to update the text box with the word list
def update_text_box(word_list):
    text_box.delete("1.0", tk.END)
    for word in word_list:
        text_box.insert(tk.END, word + "\n")

# Function to save the word list to a text file
def save_word_list():
    global saved_word_list
    file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt")])
    if file_path:
        with open(file_path, "w") as file:
            for word in saved_word_list:
                file.write(word + "\n")

# Function to load the word list from a text file
def load_word_list():
    global saved_word_list
    file_path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt")])
    if file_path:
        with open(file_path, "r") as file:
            saved_word_list = [line.strip() for line in file]
            update_text_box(saved_word_list)  # Update the text box with loaded word list

# Create the main application window
app = tk.Tk()
app.title("Word Document Processor")

# Create and configure widgets
text_box = tk.Text(app, height=10, width=40, font=("Arial", 12))
text_box.pack(padx=20, pady=10)

button_frame = tk.Frame(app)
button_frame.pack(pady=20)

update_word_list_button = tk.Button(button_frame, text="Update Word List", font=("Arial", 14), command=update_word_list)
update_word_list_button.grid(row=0, column=0, padx=5)

save_word_list_button = tk.Button(button_frame, text="Save Word List", font=("Arial", 14), command=save_word_list)
save_word_list_button.grid(row=0, column=1, padx=5)

load_word_list_button = tk.Button(button_frame, text="Load Word List", font=("Arial", 14), command=load_word_list)
load_word_list_button.grid(row=0, column=2, padx=5)

process_button = tk.Button(app, text="Process", font=("Arial", 14), command=process_document)
process_button.pack(pady=20)

# Start the GUI application
app.mainloop()
