import docx
import os

def read_word_list_from_docx(docx_filename):
    try:
        doc = docx.Document(docx_filename)
        word_list = [paragraph.text.strip() for paragraph in doc.paragraphs]
        return word_list
    except FileNotFoundError:
        print(f"Error: The file '{docx_filename}' was not found.")
        return []

def bracket_matching_words(input_text, word_list):
    words = input_text.split()
    result = []

    for word in words:
        # Check if the word matches any word in the provided list (case-insensitive)
        if any(word.lower() == item.lower() for item in word_list):
            result.append(f"[{word}]")
        else:
            result.append(word)

    result_text = ' '.join(result)
    return result_text

if __name__ == "__main__":
    while True:
        # Prompt the user to input the path to the Word document-based word list
        word_list_path = input("Enter the path to the Word document-based word list: ").strip()

        # Read the word list from the Word document
        word_list = read_word_list_from_docx(word_list_path)

        if not word_list:
            print("Word list is empty or could not be read.")
            continue

        # Specify the path to the input Word document
        input_docx_path = r"C:\Users\HamburgerC\Desktop\Bracketing code\Unbracketed.docx"

        # Create a new Word document to store the result
        result_doc = docx.Document()

        # Load the input Word document
        try:
            doc = docx.Document(input_docx_path)
        except FileNotFoundError:
            print(f"Error: The file '{input_docx_path}' was not found.")
            break

        # Initialize an empty list to store the processed paragraphs
        processed_paragraphs = []

        # Iterate through the paragraphs in the document
        for paragraph in doc.paragraphs:
            # Process each paragraph's text and add it to the list
            processed_text = bracket_matching_words(paragraph.text, word_list)
            processed_paragraphs.append(processed_text)

        # Add processed paragraphs to the result document
        for processed_text in processed_paragraphs:
            result_doc.add_paragraph(processed_text)

        # Specify the path for the output Word document
        output_docx_path = r"C:\Users\HamburgerC\Desktop\Bracketing code\Processed_Bracketed.docx"

        # Save the result to a new Word document
        try:
            result_doc.save(output_docx_path)
            print(f"Processing successful. Processed document saved as '{output_docx_path}'.")
        except Exception as e:
            print(f"Error: An error occurred while saving the processed document: {str(e)}")
            print("Processing failed.")

        # Ask the user if they want to process another word list
        another_word_list = input("Do you want to process another word list? (yes/no): ").strip().lower()
        if another_word_list != 'yes':
            break
