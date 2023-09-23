import docx
import os
import logging
import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
from tkinter import simpledialog

logging.basicConfig(level=logging.ERROR)

def add_double_brackets_to_words_in_text(input_text, word_list):
    words = input_text.split()
    result = []

    for word in words:
        if any(word.lower() == item.lower() for item in word_list):
            result.append(f"[[{word}]]")
        else:
            result.append(word)

    result_text = ' '.join(result)
    return result_text

def extract_text_from_docx(docx_filename):
    try:
        doc = docx.Document(docx_filename)
        text = '\n'.join([paragraph.text.strip() for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        logging.error(f"Error: Unable to read the Word document. Error: {str(e)}")
        return ""

def process_word_document(docx_filename, word_list, result_filename):
    try:
        doc_text = extract_text_from_docx(docx_filename)
    except Exception as e:
        logging.error(f"Error: Unable to read the Word document. Error: {str(e)}")
        return False

    paragraphs = doc_text.split('\n')
    processed_paragraphs = []

    for paragraph in paragraphs:
        processed_paragraph = add_double_brackets_to_words_in_text(paragraph, word_list)
        processed_paragraphs.append(processed_paragraph)

    result_text = '\n'.join(processed_paragraphs)

    full_output_path = os.path.join(result_filename)

    if os.path.exists(full_output_path):
        confirm = messagebox.askyesno("File Exists", f"The file '{full_output_path}' already exists. Do you want to overwrite it?")
        if not confirm:
            print("Processing canceled.")
            return False

    result_doc = docx.Document()
    result_doc.add_paragraph(result_text)

    try:
        result_doc.save(full_output_path)
        return True
    except Exception as e:
        logging.error(f"Error: Unable to save the Word document. Error: {str(e)}")
        return False

def update_word_list(word_list, text_box):
    word_list.clear()
    word_list.extend(text_box.get("1.0", "end-1c").splitlines())

def add_word(word_list, text_box):
    new_word = simpledialog.askstring("Add Word", "Enter a word to add:")
    if new_word:
        word_list.append(new_word)
        update_text_box(word_list, text_box)

def remove_word(word_list, text_box):
    word_to_remove = simpledialog.askstring("Remove Word", "Enter a word to remove:")
    if word_to_remove in word_list:
        word_list.remove(word_to_remove)
        update_text_box(word_list, text_box)

def update_text_box(word_list, text_box):
    text_box.delete("1.0", tk.END)
    for word in word_list:
        text_box.insert(tk.END, word + "\n")

def start_application():
    app = tk.Tk()
    app.title("Word Document Processor")
    app.configure(bg="#808080")

    label = tk.Label(app, text="Word List:", fg="white", bg="#808080", font=("Arial", 14))
    text_box = tk.Text(app, height=10, width=40, font=("Arial", 12))
    update_word_list_button = tk.Button(app, text="Update Word List", fg="white", bg="#007AFF", font=("Arial", 14), command=lambda: update_word_list(saved_word_list, text_box))
    
    # Define the process_document function here before using it as a command
    def process_document():
        global saved_word_list
        word_list = text_box.get("1.0", "end-1c").splitlines()
        if not word_list:
            messagebox.showerror("Error", "Word list is empty.")
            return

        output_filename = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not output_filename:
            return

        if process_word_document(output_filename, word_list, output_filename):
            messagebox.showinfo("Success", f"Processing completed successfully. Result saved as '{output_filename}'")
        else:
            messagebox.showerror("Error", "Processing failed.")
    
    process_button = tk.Button(app, text="Process", fg="white", bg="#007AFF", font=("Arial", 14), command=process_document)
    add_word_button = tk.Button(app, text="Add Word", fg="white", bg="#007AFF", font=("Arial", 14), command=lambda: add_word(saved_word_list, text_box))
    remove_word_button = tk.Button(app, text="Remove Word", fg="white", bg="#007AFF", font=("Arial", 14), command=lambda: remove_word(saved_word_list, text_box))

    label.pack(pady=20)
    text_box.pack(padx=20, pady=10)
    update_word_list_button.pack(pady=10)
    process_button.pack(pady=10)
    add_word_button.pack(pady=10)
    remove_word_button.pack(pady=10)

    app.mainloop()

saved_word_list = []

if __name__ == "__main__":
    start_application()
